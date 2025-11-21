"""
Pytest configuration and fixtures for the grading app tests.
"""

import os
import sys
import tempfile
from datetime import datetime
from unittest.mock import MagicMock, patch


def pytest_configure(config):
    """Configure pytest before test collection.

    Ensures sys.path is set before any test modules are imported
    across all pytest-xdist workers.
    """
    from pathlib import Path
    # Add project root to Python path dynamically
    project_root = Path(__file__).parent.parent
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))


import pytest

from services.deployment_service import DeploymentService

# Mock desktop-specific dependencies before importing app
# This is needed because app.py imports routes which import desktop modules
sys.modules["webview"] = MagicMock()
sys.modules["keyring"] = MagicMock()
sys.modules["keyrings"] = MagicMock()
sys.modules["keyrings.cryptfile"] = MagicMock()
sys.modules["keyrings.cryptfile.cryptfile"] = MagicMock()
sys.modules["apscheduler"] = MagicMock()
sys.modules["apscheduler.schedulers"] = MagicMock()
sys.modules["apscheduler.schedulers.background"] = MagicMock()
sys.modules["pystray"] = MagicMock()
sys.modules["PIL"] = MagicMock()
sys.modules["PIL.Image"] = MagicMock()
# Note: Desktop modules are mocked in individual test files as needed
# to avoid interfering with desktop-specific tests


# Mock Redis for password reset token tests
class MockRedis:
    def __init__(self):
        self.data = {}
        self.ttl_data = {}

    def get(self, key):
        if key in self.ttl_data and datetime.now().timestamp() > self.ttl_data[key]:
            self.data.pop(key, None)
            self.ttl_data.pop(key, None)
            return None
        return self.data.get(key)

    def set(self, key, value, ex=None):
        self.data[key] = value
        if ex:
            self.ttl_data[key] = datetime.now().timestamp() + ex

    def setex(self, key, time, value):
        self.data[key] = value
        self.ttl_data[key] = datetime.now().timestamp() + time

    def delete(self, key):
        self.data.pop(key, None)
        self.ttl_data.pop(key, None)

    @staticmethod
    def from_url(url, decode_responses=True):
        return MockRedis()


# Create RedisError exception for testing
class RedisError(Exception):
    pass


# Mock redis module
sys.modules["redis"] = MagicMock()
sys.modules["redis"].from_url = MockRedis.from_url
sys.modules["redis"].RedisError = RedisError


# Mock Celery task decorator to add .delay method to functions
class MockCeleryTask:
    """Mock Celery task that adds .delay and .apply_async methods."""

    def __call__(self, func):
        """Decorate function to add Celery task methods."""
        # Create a wrapper function that preserves the original function's attributes
        def wrapped_func(*args, **kwargs):
            return func(*args, **kwargs)
        
        # Copy all attributes from the original function to the wrapper
        wrapped_func.__name__ = func.__name__
        wrapped_func.__doc__ = func.__doc__
        wrapped_func.__module__ = func.__module__
        wrapped_func.__qualname__ = getattr(func, '__qualname__', func.__name__)
        
        # Add Celery task methods to the wrapper
        mock_delay = MagicMock(return_value=MagicMock(id="mock-task-id"))
        mock_apply_async = MagicMock(return_value=MagicMock(id="mock-task-id"))
        wrapped_func.delay = mock_delay
        wrapped_func.apply_async = mock_apply_async
        
        return wrapped_func


# Patch tasks module to add .delay to all task functions
def mock_tasks_module():
    """Add .delay method to task functions after tasks module is imported."""
    try:
        import tasks

        # Add .delay method to process_job and other task functions
        if hasattr(tasks, "process_job") and not hasattr(tasks.process_job, "delay"):
            tasks.process_job.delay = MagicMock(
                return_value=MagicMock(id="mock-task-id")
            )
            tasks.process_job.apply_async = MagicMock(
                return_value=MagicMock(id="mock-task-id")
            )
    except ImportError:
        pass

# Set TESTING environment variable and override DATABASE_URL before importing app
# This ensures tests use SQLite instead of PostgreSQL
os.environ["TESTING"] = "True"
os.environ["DATABASE_URL"] = (
    "sqlite:///:memory:"  # Override to use in-memory SQLite for tests
)

# Import the app and models
from app import app as flask_app
from models import GradingJob, JobBatch, MarkingScheme, Submission, db

# Mock task functions to have .delay method
mock_tasks_module()


@pytest.fixture(scope="session", autouse=True)
def initialize_db_for_tests():
    """
    Session-level fixture to ensure db is properly initialized with the app.
    This runs once at the start of the test session and ensures that the
    db instance is properly bound to the Flask app, preventing SQLAlchemy
    "not registered" errors.
    """
    # Force re-initialization of db with the app if needed
    if "sqlalchemy" not in flask_app.extensions:
        db.init_app(flask_app)
    yield


@pytest.fixture
def app():
    """Create a Flask app for testing."""
    # Use in-memory database for tests to avoid concurrency issues
    # Each test gets its own isolated in-memory database

    # Create a temporary upload folder for this test
    upload_folder = tempfile.mkdtemp()

    # Configure the app for testing
    flask_app.config.update(
        {
            "TESTING": True,
            "SQLALCHEMY_DATABASE_URI": "sqlite:///:memory:",
            "SQLALCHEMY_TRACK_MODIFICATIONS": False,
            "WTF_CSRF_ENABLED": False,
            "UPLOAD_FOLDER": upload_folder,
            "RATELIMIT_ENABLED": False,
            "SQLALCHEMY_ENGINE_OPTIONS": {
                "pool_pre_ping": True,
                "pool_recycle": 300,
                "connect_args": {
                    "check_same_thread": False,
                    "timeout": 30,
                },
            },
        }
    )

    # Disable rate limiting for tests by monkey-patching the limiter
    try:
        from app import limiter

        if hasattr(limiter, "_enabled"):
            limiter._enabled = False
        elif hasattr(limiter, "enabled"):
            limiter.enabled = False
    except (ImportError, AttributeError):
        pass

    # Ensure db is initialized with app
    # This fixes "The current Flask app is not registered with this
    # 'SQLAlchemy'" error which can happen if the app was imported before
    # db was fully ready or in some test environments
    if "sqlalchemy" not in flask_app.extensions:
        db.init_app(flask_app)

    # Register this app instance with tasks module so worker threads use it
    import tasks
    tasks.set_test_app(flask_app)

    # Create the database tables
    with flask_app.app_context():
        db.create_all()
        yield flask_app

        # Clear the registered test app
        tasks.clear_test_app()

        # Cleanup: ensure all sessions are closed and removed
        try:
            db.session.rollback()
        except Exception:
            pass
        try:
            db.session.close()
        except Exception:
            pass
        try:
            db.session.remove()
        except Exception:
            pass
        try:
            db.drop_all()
        except Exception:
            pass
        # Close all connections
        try:
            db.engine.dispose()
        except Exception:
            pass

    # Clean up the temporary upload folder
    try:
        import shutil
        shutil.rmtree(upload_folder, ignore_errors=True)
    except Exception:
        pass


@pytest.fixture(autouse=True)
def cleanup_db_session(app):
    """Automatically clean up database session after each test."""
    yield
    # Clean up after each test
    with app.app_context():
        try:
            db.session.rollback()
        except Exception:
            pass
        try:
            db.session.close()
        except Exception:
            pass
        try:
            db.session.remove()
        except Exception:
            pass


@pytest.fixture(autouse=True)
def reset_deployment_mode(app):
    """Reset deployment mode to multi-user after each test by default.

    Tests that specifically need single-user mode use dedicated
    fixtures (e.g. app_single_user), so the default for generic
    tests can be multi-user to exercise auth middleware.
    """
    yield
    # Reset to multi-user mode after each test
    with app.app_context():
        try:
            from models import DeploymentConfig
            from services.deployment_service import DeploymentService

            # Check if deployment_config table exists before trying to reset
            # This handles tests that create their own isolated database
            inspector = db.inspect(db.engine)
            if "deployment_config" in inspector.get_table_names():
                DeploymentService.set_mode("multi-user")
        except Exception:
            # Silently ignore errors - test may have its own isolated DB
            pass


@pytest.fixture(autouse=True)
def cleanup_flask_globals():
    """Clean up Flask globals and request context after each test."""
    yield
    # Clean up Flask globals
    try:
        from flask import g

        # Clear Flask g object
        if hasattr(g, "__dict__"):
            g.__dict__.clear()
    except Exception:
        pass


@pytest.fixture(autouse=True)
def cleanup_scheduler():
    """Clean up scheduler state before and after tests."""
    import time

    # Clean up BEFORE test to ensure fresh state
    try:
        import desktop.scheduler

        # Shutdown the global scheduler if it's running
        if hasattr(desktop.scheduler, 'scheduler'):
            scheduler = desktop.scheduler.scheduler
            if hasattr(scheduler, 'running') and scheduler.running:
                scheduler.shutdown(wait=True)
                # Give it a moment to fully stop and release all resources
                time.sleep(0.3)

            # Remove all jobs to ensure clean state
            if hasattr(scheduler, 'remove_all_jobs'):
                try:
                    scheduler.remove_all_jobs()
                except Exception:
                    pass

            # Create a fresh scheduler instance for the next test
            from apscheduler.schedulers.background import BackgroundScheduler
            desktop.scheduler.scheduler = BackgroundScheduler()
    except Exception:
        pass

    yield

    # Clean up AFTER test as well
    try:
        import desktop.scheduler

        # Shutdown the global scheduler if it's running
        if hasattr(desktop.scheduler, 'scheduler'):
            scheduler = desktop.scheduler.scheduler
            if hasattr(scheduler, 'running') and scheduler.running:
                scheduler.shutdown(wait=True)
                # Give it a moment to fully stop and release all resources
                time.sleep(0.3)

            # Remove all jobs to ensure clean state
            if hasattr(scheduler, 'remove_all_jobs'):
                try:
                    scheduler.remove_all_jobs()
                except Exception:
                    pass

            # Create a fresh scheduler instance for the next test
            from apscheduler.schedulers.background import BackgroundScheduler
            desktop.scheduler.scheduler = BackgroundScheduler()
    except Exception:
        pass


@pytest.fixture(autouse=True)
def cleanup_environment():
    """Clean up test-specific environment variables after each test."""
    # Store original test-related env vars
    test_env_vars = [
        "FLASK_ENV",
        "FLASK_DEBUG",
        "SECRET_KEY",
        "DB_ENCRYPTION_KEY",
        "OPENROUTER_API_KEY",
        "CLAUDE_API_KEY",
        "LM_STUDIO_URL",
    ]
    original_values = {key: os.environ.get(key) for key in test_env_vars}
    yield
    # Restore original values for test-related env vars only
    for key in test_env_vars:
        if original_values[key] is not None:
            os.environ[key] = original_values[key]
        elif key in os.environ:
            del os.environ[key]


@pytest.fixture
def client(app):
    """Create a test client for the Flask app."""
    return app.test_client()


@pytest.fixture
def runner(app):
    """Create a test runner for the Flask app."""
    return app.test_cli_runner()


@pytest.fixture
def sample_job(app):
    """Create a sample grading job for testing."""
    with app.app_context():
        from models import db

        job = GradingJob(
            job_name="Test Job",
            description="A test job for unit testing",
            provider="openrouter",
            model="anthropic/claude-3-5-sonnet-20241022",
            prompt="Please grade this document.",
            priority=5,
            temperature=0.7,
            max_tokens=2000,
        )
        db.session.add(job)
        db.session.commit()

        # Store the ID and return a fresh object each time it's accessed

        # Store the ID for later retrieval
        return job


@pytest.fixture
def sample_submission(app, sample_job):
    """Create a sample submission for testing."""
    with app.app_context():
        from models import db

        submission = Submission(
            job_id=sample_job.id,
            original_filename="test_document.txt",
            filename="test_document.txt",
            file_type="txt",
            file_size=1024,
            status="failed",
        )
        db.session.add(submission)
        db.session.commit()

        # Store the ID for later retrieval
        return submission


@pytest.fixture
def sample_marking_scheme(app):
    """Create a sample marking scheme for testing."""
    with app.app_context():
        from models import db

        marking_scheme = MarkingScheme(
            name="Test Marking Scheme",
            original_filename="test_rubric.txt",
            filename="test_rubric.txt",
            file_type="txt",
            file_size=2048,
            content="Test marking scheme content",
        )
        db.session.add(marking_scheme)
        db.session.commit()
        db.session.refresh(marking_scheme)
        return marking_scheme


@pytest.fixture
def sample_batch(app):
    """Create a sample batch for testing."""
    with app.app_context():
        from models import db

        batch = JobBatch(
            batch_name="Test Batch",
            description="A test batch for unit testing",
            status="pending",
            priority=5,
        )
        db.session.add(batch)
        db.session.commit()

        # Store the ID and return a fresh object each time it's accessed
        batch_id = batch.id

        # Clear the session to ensure fresh object retrieval in tests
        db.session.expunge(batch)

        # Return a fresh instance
        fresh_batch = db.session.get(JobBatch, batch_id)
        return fresh_batch


@pytest.fixture
def mock_api_keys():
    """Mock API keys for testing."""
    with patch.dict(
        os.environ,
        {
            "OPENROUTER_API_KEY": "test-openrouter-key",
            "CLAUDE_API_KEY": "test-claude-key",
            "LM_STUDIO_URL": "http://localhost:1234/v1",
        },
    ):
        yield


@pytest.fixture
def mock_celery():
    """Mock Celery for testing."""
    with patch("tasks.celery_app") as mock_celery_app:
        mock_celery_app.send_task.return_value = MagicMock(id="test-task-id")
        yield mock_celery_app


@pytest.fixture
def sample_text_file():
    """Create a sample text file for testing."""
    content = (
        "This is a test document for grading. It contains sample text to evaluate."
    )
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write(content)
        temp_path = f.name

    yield temp_path

    # Clean up
    os.unlink(temp_path)


@pytest.fixture
def sample_docx_file():
    """Create a sample DOCX file for testing."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("This is a test document for grading.")
    doc.add_paragraph("It contains multiple paragraphs to evaluate.")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        temp_path = f.name

    yield temp_path

    # Clean up
    os.unlink(temp_path)


@pytest.fixture
def sample_pdf_file():
    """Create a sample PDF file for testing."""

    from PyPDF2 import PdfWriter

    # Create a simple PDF
    PdfWriter()
    # Note: This is a simplified PDF creation - in real tests you might want to use a pre-made PDF

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        # For testing purposes, we'll create an empty PDF file
        pdf_content = (
            b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n1 0 obj\n<<\n/Type /Catalog\n"
            b"/Pages 2 0 R\n>>\nendobj\n2 0 obj\n<<\n/Type /Pages\n/Kids []\n"
            b"/Count 0\n>>\nendobj\nxref\n0 3\n0000000000 65535 f \n"
            b"0000000009 00000 n \n0000000058 00000 n \ntrailer\n<<\n/Size 3\n"
            b"/Root 1 0 R\n>>\nstartxref\n108\n%%EOF\n"
        )
        f.write(pdf_content)
        temp_path = f.name

    yield temp_path

    # Clean up
    os.unlink(temp_path)


@pytest.fixture
def multi_user_mode(app):
    """Set deployment to multi-user mode."""
    from services.deployment_service import DeploymentService

    with app.app_context():
        DeploymentService.set_mode("multi-user")
        yield
        # Cleanup: reset to single-user mode
        DeploymentService.set_mode("single-user")


@pytest.fixture
def test_user(app, multi_user_mode):
    """Create a regular test user for authentication tests."""
    from services.auth_service import AuthService

    with app.app_context():
        user = AuthService.create_user(
            "testuser@example.com", "TestPass123!", "Test User", is_admin=False
        )
        yield user


@pytest.fixture
def admin_user(app, multi_user_mode):
    """Create an admin user for authentication tests."""
    from services.auth_service import AuthService

    with app.app_context():
        user = AuthService.create_user(
            "admin@example.com", "AdminPass123!", "Admin User", is_admin=True
        )
        yield user


@pytest.fixture
def auth_headers(client, test_user):
    """Get authentication headers for a regular user."""
    response = client.post(
        "/api/auth/login",
        json={"email": "testuser@example.com", "password": "TestPass123!"},
    )
    assert response.status_code == 200
    # Flask-Login uses session cookies, so no need for explicit headers
    # The session is maintained by the test client
    return {}


@pytest.fixture
def admin_headers(client, admin_user):
    """Get authentication headers for an admin user."""
    response = client.post(
        "/api/auth/login",
        json={"email": "admin@example.com", "password": "AdminPass123!"},
    )
    assert response.status_code == 200
    # Flask-Login uses session cookies, so no need for explicit headers
    # The session is maintained by the test client
    return {}


class AuthActions:
    """Helper class for authentication actions in tests."""

    def __init__(self, client, app):
        self._client = client
        self._app = app

    def login(self, email, password):
        """Log in a user."""
        return self._client.post(
            "/api/auth/login",
            json={"email": email, "password": password},
        )

    def logout(self):
        """Log out the current user."""
        return self._client.post("/api/auth/logout")


@pytest.fixture
def auth(client, app):
    """Provide authentication helper for tests."""
    return AuthActions(client, app)


@pytest.fixture
def app_multi_user(app):
    """Configure shared test app in multi-user mode."""
    with app.app_context():
        DeploymentService.set_mode("multi-user")
    return app


@pytest.fixture
def client_single_user(app_single_user):
    """Create test client for single-user mode using shared app fixture."""
    return app_single_user.test_client()


@pytest.fixture
def client_multi_user(app_multi_user):
    """Create test client for multi-user mode using shared app fixture."""
    return app_multi_user.test_client()


@pytest.fixture
def test_project(app, test_user):
    """Create a test project/grading job owned by test_user."""
    with app.app_context():
        from models import db

        job = GradingJob(
            job_name="Test Project",
            description="A test project for testing",
            provider="openrouter",
            model="anthropic/claude-3-5-sonnet-20241022",
            prompt="Please grade this document.",
            owner_id=test_user.id,
            priority=5,
            temperature=0.7,
            max_tokens=2000,
        )
        db.session.add(job)
        db.session.commit()
        db.session.refresh(job)
        return job
