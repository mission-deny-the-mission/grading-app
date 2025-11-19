"""
Sample Documents for Testing

Provides mock PDF, DOCX, and image files for document parsing tests.
"""

import os
import tempfile
from pathlib import Path

# Try to import optional libraries for creating real test files
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    from PIL import Image
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False


def create_sample_pdf(content: str = None) -> str:
    """
    Create a temporary PDF file for testing.

    Args:
        content: Text content to include in PDF

    Returns:
        str: Path to temporary PDF file
    """
    temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    temp_file.close()

    if HAS_REPORTLAB:
        try:
            c = canvas.Canvas(temp_file.name, pagesize=letter)
            text = content or "Sample PDF Document\n\nThis is a test document."
            y_position = 750
            for line in text.split('\n'):
                c.drawString(50, y_position, line)
                y_position -= 20
            c.save()
            return temp_file.name
        except Exception:
            pass

    # Fallback: create a text file with PDF-like content
    with open(temp_file.name, 'w') as f:
        f.write(content or "Sample PDF Document\n\nThis is a test document.")
    return temp_file.name


def create_sample_docx(content: str = None) -> str:
    """
    Create a temporary DOCX file for testing.

    Args:
        content: Text content to include in DOCX

    Returns:
        str: Path to temporary DOCX file
    """
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    temp_file.close()

    if HAS_PYTHON_DOCX:
        try:
            doc = DocxDocument()
            text = content or "Sample DOCX Document\n\nThis is a test document."
            for line in text.split('\n'):
                doc.add_paragraph(line)
            doc.save(temp_file.name)
            return temp_file.name
        except Exception:
            pass

    # Fallback: write minimal DOCX XML structure
    with open(temp_file.name, 'wb') as f:
        # Minimal DOCX is a ZIP file, but for testing we just need something that
        # looks like a DOCX file
        f.write(b'PK\x03\x04')  # ZIP file header
    return temp_file.name


def create_sample_png() -> str:
    """
    Create a temporary PNG image file for testing.

    Returns:
        str: Path to temporary PNG file
    """
    temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    temp_file.close()

    if HAS_PILLOW:
        try:
            img = Image.new('RGB', (100, 100), color='white')
            img.save(temp_file.name, 'PNG')
            return temp_file.name
        except Exception:
            pass

    # Fallback: write minimal PNG header
    with open(temp_file.name, 'wb') as f:
        f.write(b'\x89PNG\r\n\x1a\n')
    return temp_file.name


def create_sample_jpg() -> str:
    """
    Create a temporary JPG image file for testing.

    Returns:
        str: Path to temporary JPG file
    """
    temp_file = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
    temp_file.close()

    if HAS_PILLOW:
        try:
            img = Image.new('RGB', (100, 100), color='white')
            img.save(temp_file.name, 'JPEG')
            return temp_file.name
        except Exception:
            pass

    # Fallback: write minimal JPG header
    with open(temp_file.name, 'wb') as f:
        f.write(b'\xff\xd8\xff\xe0')
    return temp_file.name


def create_sample_rubric_pdf() -> str:
    """
    Create a PDF with rubric content for LLM testing.

    Returns:
        str: Path to temporary PDF file with rubric text
    """
    rubric_content = """ESSAY GRADING RUBRIC

Criterion 1: Organization (20 points)
- Excellent (20): Clear introduction, body, and conclusion with logical flow
- Good (15): Most sections present with mostly logical flow
- Satisfactory (10): Basic structure but flow is unclear
- Poor (5): Unclear structure or missing sections
- Fail (0): No clear structure

Criterion 2: Grammar and Mechanics (20 points)
- Excellent (20): Virtually no errors
- Good (15): Few errors that don't interfere with understanding
- Satisfactory (10): Some errors but most sentences are correct
- Poor (5): Many errors that sometimes interfere
- Fail (0): Numerous errors preventing understanding"""

    temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    temp_file.close()

    if HAS_REPORTLAB:
        try:
            c = canvas.Canvas(temp_file.name, pagesize=letter)
            y_position = 750
            for line in rubric_content.split('\n'):
                c.drawString(50, y_position, line)
                y_position -= 15
            c.save()
            return temp_file.name
        except Exception:
            pass

    # Fallback: write as text file
    with open(temp_file.name, 'w') as f:
        f.write(rubric_content)
    return temp_file.name


def create_sample_rubric_docx() -> str:
    """
    Create a DOCX with rubric content for LLM testing.

    Returns:
        str: Path to temporary DOCX file with rubric text
    """
    rubric_content = """PROJECT GRADING RUBRIC

Requirements Analysis (15 points)
Complete requirements gathering with stakeholder input and documentation - Excellent
Good requirements with minor gaps - Good
Basic requirements but some missing - Satisfactory
Incomplete or unclear requirements - Poor

System Design (20 points)
Excellent architecture with scalability - Excellent
Good design with minor improvements needed - Good
Adequate design but lacks optimization - Satisfactory
Poor design with significant issues - Poor"""

    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    temp_file.close()

    if HAS_PYTHON_DOCX:
        try:
            doc = DocxDocument()
            for line in rubric_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            doc.save(temp_file.name)
            return temp_file.name
        except Exception:
            pass

    # Fallback: write minimal DOCX structure
    with open(temp_file.name, 'wb') as f:
        f.write(b'PK\x03\x04')  # ZIP file header
    return temp_file.name


def cleanup_temp_file(file_path: str) -> None:
    """
    Clean up a temporary file.

    Args:
        file_path: Path to temporary file to delete
    """
    if file_path and os.path.exists(file_path):
        os.remove(file_path)


def get_all_sample_documents():
    """
    Get paths to all sample documents (for cleanup).

    Returns:
        list: List of temporary file paths
    """
    # TODO: Track created files for cleanup (T013)
    return []
